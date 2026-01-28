import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

# --- CẤU HÌNH API GEMINI ---
try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
except:
    API_KEY = os.getenv("GOOGLE_API_KEY", "")

if not API_KEY:
    st.error("❌ Lỗi: Chưa cấu hình API Key. Vui lòng thêm GOOGLE_API_KEY vào secrets hoặc biến môi trường.")
    st.stop()

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

# --- HÀM XỬ LÝ XUẤT FILE WORD CHUẨN ---
def create_docx(content, title):
    doc = Document()
    
    # Cấu hình định dạng trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1)

    # Style chung: Times New Roman, cỡ 13
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Tiêu đề: In đậm, viết hoa, căn giữa
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr.add_run(f"KẾ HOẠCH BÀI DẠY STEM: {title.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    # Xử lý nội dung AI trả về
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Xử lý in đậm cho các mục chính (giả định AI dùng dấu ** hoặc ##)
        clean_line = line.replace('**', '').replace('##', '').replace('#', '')
        run = p.add_run(clean_line)
        
        if line.startswith('##') or line.startswith('HĐ') or line.startswith('Hoạt động'):
            run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- GIAO DIỆN STREAMLIT ---
st.set_page_config(page_title="STEM Math 3089 Planner", layout="wide")
st.title("📐 Chương trình hỗ trợ soạn thiết kế " \
"bài dạy STEM Toán học (CV 3089)")

# Sidebar nhập liệu
with st.sidebar:
    st.header("📋 Thông số bài dạy")
    lop = st.selectbox("Khối lớp", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"])
    ten_bai = st.text_input("Tên bài dạy", placeholder="VD: Thiết kế kim tự tháp")
    chu_trinh = st.selectbox("Chu trình STEM", ["Chu trình kỹ thuật (EDP)", "Chu trình khoa học"])
    thoi_luong = st.selectbox("Thời lượng", ["1 tiết (45 phút)", "2 tiết (90 phút)", "3 tiết (135 phút)"])
    
    st.subheader("🔍 Chỉ dẫn thêm")
    kt_nen = st.text_area("Kiến thức nền (Toán học trọng tâm)", placeholder="VD: Định lý Thales, Tỉ số lượng giác...")
    san_pham = st.text_input("Sản phẩm dự kiến", placeholder="VD: Mô hình đo khoảng cách")
    yeu_cau_khac = st.text_area("Yêu cầu khác (nếu có)")

    hoat_dong_chon = st.multiselect(
        "Hoạt động cần soạn chi tiết",
        ["Tất cả", "HĐ 1", "HĐ 2", "HĐ 3", "HĐ 4", "HĐ 5"],
        default=["Tất cả"]
    )

# Màn hình chính
if st.button("🚀 Bắt đầu soạn "):
    if not ten_bai:
        st.warning("Vui lòng nhập tên bài dạy!")
    else:
        # Prompt Engineering tối ưu cho Toán học và CV 3089
        prompt = f"""
        Bạn là một chuyên gia giáo dục STEM với chuyên môn sâu về Toán học tại Việt Nam. 
        Hãy soạn một kế hoạch bài dạy STEM cho học sinh {lop} bài '{ten_bai}' theo CV 3089/BGDĐT-GDTrH.
        
        THÔNG TIN ĐẦU VÀO:
        - Thời lượng: {thoi_luong}.
        - Chu trình: {chu_trinh}.
        - Kiến thức nền trọng tâm: {kt_nen}.
        - Sản phẩm: {san_pham}.
        - Yêu cầu khác: {yeu_cau_khac}.

        CẤU TRÚC BÀI SOẠN PHẢI BAO GỒM:
        1. Mục tiêu bài học:
           - Kiến thức (Trọng tâm Toán học).
           - Kĩ năng & Phẩm chất.
           - Năng lực đặc thù: Giải quyết vấn đề, mô hình hóa toán học.
        2. Thiết bị dạy học và học liệu: (Tự động gợi ý danh sách vật liệu tái chế phù hợp).
        3. Tiến trình dạy học (5 hoạt động theo CV 3089):
           - HĐ 1: Xác định vấn đề (Giao nhiệm vụ thực tiễn).
           - HĐ 2: Nghiên cứu kiến thức nền và đề xuất giải pháp.
           - HĐ 3: Lựa chọn giải pháp/Thiết kế sản phẩm (Vẽ bản vẽ, lập kế hoạch tính toán).
           - HĐ 4: Chế tạo mẫu, thử nghiệm và thảo luận (Thực hiện tính toán và điều chỉnh).
           - HĐ 5: Chia sẻ, thảo luận và đánh giá (Thuyết trình ứng dụng toán học).
        4. Phụ lục: 
           - Phiếu học tập cho học sinh.
           - Bảng tiêu chí đánh giá sản phẩm (Rubric) chi tiết.

        YÊU CẦU TRÌNH BÀY: Chuyên nghiệp, ngôn ngữ sư phạm, trình bày rõ ràng từng mục.
        """

        with st.spinner("⏳ AI đang soạn giáo án STEM Toán học..."):
            try:
                # Gọi Gemini API
                response = model.generate_content(prompt)
                full_text = response.text
                
                st.success("✅ Đã soạn xong giáo án!")
                
                # Nút tải file Word
                docx_data = create_docx(full_text, ten_bai)
                st.download_button(
                    label="📥 Tải xuống File Word (.docx)",
                    data=docx_data,
                    file_name=f"Giao_an_STEM_{lop.replace(' ', '_')}_{ten_bai.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Hiển thị kết quả
                st.divider()
                st.markdown("### 📄 Nội dung giáo án:")
                st.markdown(full_text)
                
            except Exception as e:
                st.error(f"❌ Lỗi khi gọi AI: {e}")
                st.info("💡 Nếu gặp lỗi quota, vui lòng cập nhật API Key mới trong file .streamlit/secrets.toml")